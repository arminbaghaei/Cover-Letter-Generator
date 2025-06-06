st.set_page_config(page_title="Cover Letter Generator", layout="centered")

# 🔷 Logo and Branding
st.image("ResearchMate1.png", width=180)
st.markdown("### Developed by **Abdollah Baghaei Daemei** – [ResearchMate.org](https://www.researchmate.org)")
st.markdown("---")

import streamlit as st
from datetime import date
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_docx(name, email, affiliation, title, author, journal, submission_date, paper_type, paper_aim, novelty, signature_path=None):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Add author header
    doc.add_paragraph(name)
    doc.add_paragraph(email)
    doc.add_paragraph(affiliation)
    doc.add_paragraph("")

    # Main letter body
    doc.add_paragraph('Dear Editor,')

    para1 = doc.add_paragraph()
    para1.add_run(f'I am pleased to submit our manuscript entitled "{title}" for consideration in ').bold = False
    para1.add_run(journal).italic = True
    para1.add_run(f' as a {paper_type}. This paper is authored by {author} and addresses the following main aim: {paper_aim}')
    para1.alignment = WD_ALIGN_PARAGRAPH.LEFT

    para2 = doc.add_paragraph()
    para2.add_run("The novelty of our work lies in: ").bold = True
    para2.add_run(novelty)
    para2.alignment = WD_ALIGN_PARAGRAPH.LEFT

    para3 = doc.add_paragraph()
    para3.add_run("We believe this work will be of interest to the readership of ").bold = False
    para3.add_run(journal).italic = True
    para3.add_run(" and aligns well with the journal’s scope. We confirm that this manuscript has not been published elsewhere and is not under consideration by any other journal.")
    para3.alignment = WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph("Thank you for considering our submission. We look forward to your positive response.")
    doc.add_paragraph("Sincerely,")

    # Signature (if provided)
    if signature_path:
        doc.add_picture(signature_path, width=Inches(1.5))

    doc.add_paragraph(author)
    doc.add_paragraph(submission_date.strftime("%B %d, %Y"))

    return doc

# Streamlit App
st.set_page_config(page_title="Cover Letter Generator", layout="centered")
st.title("📄 Manuscript Cover Letter Generator")

st.write("Fill in the details below to generate a professional manuscript cover letter.")

# Author Info
name = st.text_input("Author's Full Name")
email = st.text_input("Email Address")
affiliation = st.text_input("Affiliation")

# Manuscript Info
title = st.text_input("Manuscript Title")
author = st.text_input("Corresponding Author Name")
journal = st.text_input("Journal Name")
submission_date = st.date_input("Submission Date", value=date.today())
paper_type = st.selectbox("Paper Type", ["Original Research", "Review", "Short Communication", "Case Study"])
paper_aim = st.text_area("Main Aim or Objective of the Paper")
novelty = st.text_area("Novelty or Key Contribution")

# Signature Image Upload
signature_image = st.file_uploader("Upload your signature image (PNG or JPG, optional)", type=["png", "jpg", "jpeg"])

if st.button("Generate Cover Letter"):
    signature_path = None
    if signature_image:
        signature_path = f"temp_signature.{signature_image.name.split('.')[-1]}"
        with open(signature_path, "wb") as f:
            f.write(signature_image.read())

    doc = generate_docx(name, email, affiliation, title, author, journal, submission_date, paper_type, paper_aim, novelty, signature_path)
    doc.save("cover_letter.docx")

    with open("cover_letter.docx", "rb") as f:
        st.success("✅ Cover Letter Generated")
        st.download_button("📥 Download as Word (DOCX)", f, file_name="cover_letter.docx")
